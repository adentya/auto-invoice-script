from datetime import datetime
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Inches

import qrcode

from dotenv import load_dotenv
import emailer


load_dotenv()

document = Document(os.getenv('TEMPLATE_FILE'))

# INVOICE NO:
first_table = document.tables[0]
invoice_no_obj = first_table.cell(0, 1)
current_date = datetime.now().strftime('%m%y')
invoice_no_obj.paragraphs[0].text = None
invoice_code = os.getenv('BANK_ID_CARD')[:5]
company_code = os.getenv('YOUR_COMPANY')[:3].upper()
invoice_no = f'{company_code}{invoice_code}{current_date}'
run = invoice_no_obj.paragraphs[0].add_run(invoice_no)
run.font.color.rgb = RGBColor.from_string('9E9E9E')

# INVOICE TO:
second_table = document.tables[1]
invoice_to_obj = second_table.cell(0, 0)
invoice_to_obj.paragraphs[1].text = None
invoice_to_obj.paragraphs[1].add_run(os.getenv('YOUR_COMPANY')).bold = True

# INVOICE DATE:
second_table = document.tables[1]
invoice_date_obj = second_table.cell(0, 2)
current_date = datetime.now().strftime('%d %b %Y')
invoice_date_obj.paragraphs[1].text = None
invoice_date_obj.paragraphs[1].add_run(current_date).bold = True

# INVOICE ITEM:
third_table = document.tables[2]
invoice_item_obj = third_table.cell(1, 0)
current_date = datetime.now().strftime('%b %Y')
invoice_item_obj.text = invoice_item_obj.text.replace('[mm] [yyyy]', current_date)

# INVOICE ITEM TOTAL:
invoice_item_total_obj = third_table.cell(1, 1)
invoice_item_total_obj.paragraphs[0].text = os.getenv('YOUR_IDR_SALARY')
invoice_total_obj = third_table.cell(3, 1)
invoice_total_obj.paragraphs[0].text = None
invoice_total_obj.paragraphs[0].add_run(os.getenv('YOUR_IDR_SALARY')).bold = True

# BANK DETAILS:
bank_name_obj = document.paragraphs[10]
bank_name = bank_name_obj.text[0:bank_name_obj.text.index(':') + 2]
bank_name_obj.text = bank_name + os.getenv('BANK_ACCOUNT_NAME')

bank_account_no_obj = document.paragraphs[11]
bank_account_no = bank_account_no_obj.text[0:bank_account_no_obj.text.index(':') + 2]
bank_account_no_obj.text = bank_account_no + os.getenv('BANK_ACCOUNT_NO')

bank_full_name_obj = document.paragraphs[12]
bank_full_name = bank_full_name_obj.text[0:bank_full_name_obj.text.index(':') + 2]
bank_full_name_obj.text = bank_full_name + os.getenv('BANK_FULL_NAME')

bank_birthday_obj = document.paragraphs[13]
bank_birthday = bank_birthday_obj.text[0:bank_birthday_obj.text.index(':') + 2]
bank_birthday_obj.text = bank_birthday + os.getenv('BANK_ACCOUNT_BIRTHDAY')

bank_id_obj = document.paragraphs[14]
bank_id = bank_id_obj.text[0:bank_id_obj.text.index(':') + 2]
bank_id_obj.text = bank_id + os.getenv('BANK_ID_CARD')

# SIGNATURE IMAGE
signature_image = 'signature.png'
if os.getenv('YOUR_SIGN_IMAGE') and len(os.getenv('YOUR_SIGN_IMAGE')) \
    and os.path.isfile(os.getenv('YOUR_SIGN_IMAGE')):
    signature_image = os.getenv('YOUR_SIGN_IMAGE')
else:
    qr = qrcode.make(os.getenv('YOUR_SIGN_NAME'))
    qr.save(signature_image)
signature_img_obj = document.paragraphs[26].add_run()
signature_img_obj.add_picture(signature_image, height=Inches(1.0))

# SIGNATURE
signature_name_obj = document.paragraphs[28].text = os.getenv('YOUR_SIGN_NAME')

# GENERATE FILE
filename = 'Invoice_' + datetime.now().strftime('%b_%Y') + '.docx'
document.save(f'./generated/{filename}')

# SEND EMAIL
emailer.send_email(
    invoice_no=invoice_no, 
    month=current_date, 
    filename=filename, 
    to=os.getenv('EMAIL_TO'),
    sender_name=os.getenv('EMAIL_SENDER_NAME')
)